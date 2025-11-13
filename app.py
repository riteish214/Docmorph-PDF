import os
import secrets
import time
from datetime import datetime, timedelta
from pathlib import Path
from flask import Flask, render_template, request, send_file, jsonify, url_for
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter
from docx import Document
from pptx import Presentation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
import shutil
from utils.code_manager import generate_code, register_file, register_text, get_file_info, delete_code, cleanup_expired_codes

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SESSION_SECRET', 'dev-secret-key-change-in-production')
app.config['MAX_CONTENT_LENGTH'] = 1024 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['SHARED_FOLDER'] = 'static/shared'

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'ppt', 'pptx', 'txt'}
shared_links = {}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_old_files():
    upload_folder = Path(app.config['UPLOAD_FOLDER'])
    shared_folder = Path(app.config['SHARED_FOLDER'])
    current_time = time.time()
    
    for folder in [upload_folder, shared_folder]:
        if folder.exists():
            for file_path in folder.iterdir():
                if file_path.is_file() and file_path.name != '.gitkeep':
                    if current_time - file_path.stat().st_mtime > 3600:
                        try:
                            file_path.unlink()
                        except:
                            pass

@app.route('/')
def index():
    cleanup_old_files()
    return render_template('index.html')

@app.route('/merge', methods=['GET', 'POST'])
def merge_pdfs():
    if request.method == 'POST':
        try:
            files = request.files.getlist('files')
            if len(files) < 2:
                return jsonify({'error': 'Please upload at least 2 PDF files'}), 400
            
            pdf_writer = PdfWriter()
            
            for file in files:
                if file and file.filename and allowed_file(file.filename):
                    pdf_reader = PdfReader(file.stream)
                    for page in pdf_reader.pages:
                        pdf_writer.add_page(page)
            
            output_filename = f'merged_{int(time.time())}.pdf'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return jsonify({
                'success': True,
                'download_url': url_for('download_file', filename=output_filename)
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return render_template('merge.html')

@app.route('/split', methods=['GET', 'POST'])
def split_pdf():
    if request.method == 'POST':
        try:
            file = request.files.get('file')
            page_ranges = request.form.get('pages', '')
            
            if not file or not file.filename or not allowed_file(file.filename):
                return jsonify({'error': 'Please upload a valid PDF file'}), 400
            
            pdf_reader = PdfReader(file.stream)
            total_pages = len(pdf_reader.pages)
            
            if not page_ranges:
                page_ranges = f'1-{total_pages}'
            
            pdf_writer = PdfWriter()
            
            for page_range in page_ranges.split(','):
                page_range = page_range.strip()
                if '-' in page_range:
                    start, end = map(int, page_range.split('-'))
                    for i in range(start - 1, min(end, total_pages)):
                        pdf_writer.add_page(pdf_reader.pages[i])
                else:
                    page_num = int(page_range)
                    if 1 <= page_num <= total_pages:
                        pdf_writer.add_page(pdf_reader.pages[page_num - 1])
            
            output_filename = f'split_{int(time.time())}.pdf'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return jsonify({
                'success': True,
                'download_url': url_for('download_file', filename=output_filename),
                'total_pages': total_pages
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return render_template('split.html')

@app.route('/convert', methods=['GET', 'POST'])
def convert_file():
    if request.method == 'POST':
        temp_input = None
        try:
            file = request.files.get('file')
            convert_to = request.form.get('convert_to', 'pdf')
            
            if not file or not file.filename:
                return jsonify({'error': 'Please upload a file'}), 400
            
            filename = secure_filename(file.filename)
            if '.' not in filename:
                return jsonify({'error': 'Invalid file name'}), 400
                
            file_ext = filename.rsplit('.', 1)[1].lower()
            temp_input = os.path.join(app.config['UPLOAD_FOLDER'], f'temp_{int(time.time())}.{file_ext}')
            file.save(temp_input)
            
            output_filename = f'converted_{int(time.time())}.{convert_to}'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            if file_ext == 'pdf' and convert_to == 'docx':
                try:
                    cv = Converter(temp_input)
                    cv.convert(output_path, start=0)
                    cv.close()
                except Exception as conv_error:
                    try:
                        pdf_reader = PdfReader(temp_input)
                        doc = Document()
                        doc.add_heading('Converted from PDF', 0)
                        
                        for i, page in enumerate(pdf_reader.pages, 1):
                            extracted_text = page.extract_text()
                            if extracted_text and extracted_text.strip():
                                doc.add_heading(f'Page {i}', level=1)
                                for paragraph in extracted_text.split('\n\n'):
                                    if paragraph.strip():
                                        doc.add_paragraph(paragraph.strip())
                            else:
                                doc.add_heading(f'Page {i}', level=1)
                                doc.add_paragraph('[No extractable text on this page]')
                        
                        doc.save(output_path)
                    except Exception as fallback_error:
                        if temp_input and os.path.exists(temp_input):
                            os.remove(temp_input)
                        return jsonify({'error': f'PDF to DOCX conversion failed. Try converting to TXT format instead.'}), 500
            elif file_ext == 'pdf' and convert_to == 'txt':
                pdf_reader = PdfReader(temp_input)
                text_content = []
                for page in pdf_reader.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text_content.append(extracted_text)
                    else:
                        text_content.append('[No extractable text on this page]')
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n\n'.join(text_content))
            elif file_ext == 'docx' and convert_to == 'pdf':
                doc = Document(temp_input)
                can = canvas.Canvas(output_path, pagesize=letter)
                y_position = 750
                
                for para in doc.paragraphs:
                    text = para.text
                    if text.strip():
                        lines = text.split('\n')
                        for line in lines:
                            if line.strip():
                                can.drawString(50, y_position, line[:90])
                                y_position -= 15
                                if y_position < 50:
                                    can.showPage()
                                    y_position = 750
                
                can.save()
            elif file_ext == 'txt' and convert_to == 'pdf':
                with open(temp_input, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                
                can = canvas.Canvas(output_path, pagesize=letter)
                y_position = 750
                
                lines = text_content.split('\n')
                for line in lines:
                    if y_position < 50:
                        can.showPage()
                        y_position = 750
                    can.drawString(50, y_position, line[:90])
                    y_position -= 15
                
                can.save()
            elif file_ext == 'docx' and convert_to == 'txt':
                doc = Document(temp_input)
                text_content = []
                for para in doc.paragraphs:
                    text_content.append(para.text)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(text_content))
            elif file_ext == 'txt' and convert_to == 'docx':
                with open(temp_input, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                doc = Document()
                for line in text_content.split('\n'):
                    doc.add_paragraph(line)
                doc.save(output_path)
            else:
                if temp_input and os.path.exists(temp_input):
                    os.remove(temp_input)
                return jsonify({'error': f'Conversion from {file_ext.upper()} to {convert_to.upper()} is not supported'}), 400
            
            if temp_input and os.path.exists(temp_input):
                os.remove(temp_input)
            
            return jsonify({
                'success': True,
                'download_url': url_for('download_file', filename=output_filename)
            })
        except Exception as e:
            if temp_input and os.path.exists(temp_input):
                try:
                    os.remove(temp_input)
                except:
                    pass
            return jsonify({'error': str(e)}), 500
    
    return render_template('convert.html')

@app.route('/compress', methods=['GET', 'POST'])
def compress_pdf():
    if request.method == 'POST':
        try:
            file = request.files.get('file')
            
            if not file or not file.filename or not allowed_file(file.filename):
                return jsonify({'error': 'Please upload a valid PDF file'}), 400
            
            pdf_reader = PdfReader(file.stream)
            pdf_writer = PdfWriter()
            
            for page in pdf_reader.pages:
                page.compress_content_streams()
                pdf_writer.add_page(page)
            
            output_filename = f'compressed_{int(time.time())}.pdf'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            original_size = len(file.read())
            file.seek(0)
            compressed_size = os.path.getsize(output_path)
            reduction = ((original_size - compressed_size) / original_size) * 100
            
            return jsonify({
                'success': True,
                'download_url': url_for('download_file', filename=output_filename),
                'original_size': original_size,
                'compressed_size': compressed_size,
                'reduction': round(reduction, 2)
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return render_template('compress.html')

@app.route('/rotate', methods=['GET', 'POST'])
def rotate_pdf():
    if request.method == 'POST':
        try:
            file = request.files.get('file')
            rotation = int(request.form.get('rotation', 90))
            
            if not file or not file.filename or not allowed_file(file.filename):
                return jsonify({'error': 'Please upload a valid PDF file'}), 400
            
            pdf_reader = PdfReader(file.stream)
            pdf_writer = PdfWriter()
            
            for page in pdf_reader.pages:
                page.rotate(rotation)
                pdf_writer.add_page(page)
            
            output_filename = f'rotated_{int(time.time())}.pdf'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return jsonify({
                'success': True,
                'download_url': url_for('download_file', filename=output_filename)
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return render_template('rotate.html')

@app.route('/secure', methods=['GET', 'POST'])
def secure_pdf():
    if request.method == 'POST':
        try:
            file = request.files.get('file')
            password = request.form.get('password', '')
            
            if not file or not file.filename or not allowed_file(file.filename):
                return jsonify({'error': 'Please upload a valid PDF file'}), 400
            
            if not password:
                return jsonify({'error': 'Please provide a password'}), 400
            
            pdf_reader = PdfReader(file.stream)
            pdf_writer = PdfWriter()
            
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)
            
            pdf_writer.encrypt(password)
            
            output_filename = f'secured_{int(time.time())}.pdf'
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return jsonify({
                'success': True,
                'download_url': url_for('download_file', filename=output_filename)
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    return render_template('secure.html')

@app.route('/share', methods=['GET', 'POST'])
def share():
    if request.method == 'POST':
        try:
            share_type = request.form.get('share_type', 'file')
            
            if share_type == 'file':
                cleanup_expired_codes()
                file = request.files.get('file')
                if not file or not file.filename:
                    return jsonify({'error': 'Please upload a file'}), 400
                
                filename = secure_filename(file.filename)
                
                # Always create both link AND code
                # 1. Create link (24 hours)
                unique_id = secrets.token_urlsafe(16)
                file_path_link = os.path.join(app.config['SHARED_FOLDER'], f'{unique_id}_{filename}')
                file.save(file_path_link)
                
                share_link = url_for('get_shared', share_id=unique_id, _external=True)
                expiry = datetime.now() + timedelta(hours=24)
                
                shared_links[unique_id] = {
                    'type': 'file',
                    'path': file_path_link,
                    'filename': filename,
                    'expiry': expiry
                }
                
                # 2. Create access code (15 minutes)
                code = generate_code(6)
                codes_base = os.path.join(app.config['UPLOAD_FOLDER'], 'codes')
                code_dir = os.path.join(codes_base, code.upper())
                os.makedirs(code_dir, exist_ok=True)
                
                # Copy file for code-based access
                file_path_code = os.path.join(code_dir, filename)
                # Ensure parent directory exists
                os.makedirs(os.path.dirname(file_path_code), exist_ok=True)
                shutil.copy2(file_path_link, file_path_code)
                
                # Normalize path to absolute for consistent storage
                file_path_code_abs = os.path.normpath(os.path.abspath(file_path_code))
                
                # Register file with code manager using absolute path
                register_file(code, file_path_code_abs, filename)
                
                return jsonify({
                    'success': True,
                    'share_link': share_link,
                    'code': code,
                    'expiry': expiry.strftime('%Y-%m-%d %H:%M:%S'),
                    'filename': filename
                })
            else:
                cleanup_expired_codes()
                text_content = request.form.get('text_content', '')
                if not text_content:
                    return jsonify({'error': 'Please provide text content'}), 400
                
                # Always create both link AND code
                # 1. Create link (24 hours)
                unique_id = secrets.token_urlsafe(16)
                expiry = datetime.now() + timedelta(hours=24)
                
                shared_links[unique_id] = {
                    'type': 'text',
                    'content': text_content,
                    'expiry': expiry
                }
                
                share_link = url_for('get_shared', share_id=unique_id, _external=True)
                
                # 2. Create access code (15 minutes)
                code = generate_code(6)
                register_text(code, text_content)
                
                return jsonify({
                    'success': True,
                    'share_link': share_link,
                    'code': code,
                    'expiry': expiry.strftime('%Y-%m-%d %H:%M:%S')
                })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    
    cleanup_expired_codes()
    return render_template('share.html')

@app.route('/shared/<share_id>')
def get_shared(share_id):
    if share_id not in shared_links:
        return 'Link not found or expired', 404
    
    shared = shared_links[share_id]
    
    if datetime.now() > shared['expiry']:
        del shared_links[share_id]
        return 'Link expired', 410
    
    if shared['type'] == 'file':
        return send_file(shared['path'], as_attachment=True, download_name=shared['filename'])
    else:
        return render_template('shared_text.html', content=shared['content'])

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return 'File not found', 404

@app.route('/api/receive', methods=['POST'])
def api_receive():
    """Receive file or text by access code - integrated into share functionality"""
    try:
        cleanup_expired_codes()
        
        code = request.form.get('code', '').strip()
        if not code:
            return jsonify({'error': 'Please enter an access code'}), 400
        
        # Get info (case-insensitive)
        info = get_file_info(code)
        
        if not info:
            return jsonify({'error': 'Invalid or expired access code'}), 404
        
        info_type = info.get('type', 'file')
        
        if info_type == 'text':
            # Handle text content
            content = info.get('content', '')
            if not content:
                return jsonify({'error': 'Text content not found'}), 404
            
            # Delete code after successful retrieval
            delete_code(code)
            
            # Return text content as JSON
            return jsonify({
                'success': True,
                'type': 'text',
                'content': content
            })
        else:
            # Handle file content
            file_path = info.get('file_path')
            filename = info.get('filename')
            
            if not file_path:
                return jsonify({'error': 'File path not found'}), 404
            
            # Normalize path for Windows compatibility
            file_path = os.path.normpath(os.path.abspath(file_path))
            
            if not os.path.exists(file_path):
                # File already deleted or doesn't exist
                delete_code(code)  # Clean up the code entry
                return jsonify({'error': 'File not found or already retrieved'}), 404
            
            # Get the file content before deletion
            try:
                # Read file into memory to ensure we have it before deletion
                with open(file_path, 'rb') as f:
                    file_data = f.read()
                
                # Delete code and file after successful read
                delete_code(code)
                
                # Create a BytesIO object for sending
                file_stream = io.BytesIO(file_data)
                file_stream.seek(0)
                
                # Return file for download
                return send_file(file_stream, as_attachment=True, download_name=filename, mimetype='application/octet-stream')
            except IOError as e:
                delete_code(code)  # Clean up on error
                return jsonify({'error': f'Error reading file: {str(e)}'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['SHARED_FOLDER'], exist_ok=True)
    # Ensure codes directory exists
    codes_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'codes')
    os.makedirs(codes_dir, exist_ok=True)
    app.run(host='0.0.0.0', port=5000, debug=True)
